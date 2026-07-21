import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _md_to_para(doc: Document, text: str):
    """마크다운 줄을 Word 단락으로 변환 (헤딩, 불릿, 일반)."""
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        else:
            p = doc.add_paragraph()
            _add_inline(p, stripped)


def _add_inline(para, text: str):
    """**볼드** 처리."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def generate_proposal_word(
    project_name: str,
    rfp_result: dict,
    toc: dict,
    sections_draft: dict,
    matrix: dict,
) -> bytes:
    doc = Document()

    # ── 기본 스타일 ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # ── 표지 ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(project_name)
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run("제  안  서").font.size = Pt(18)

    doc.add_page_break()

    # ── 목차 ─────────────────────────────────────────────────────────────────
    doc.add_heading("목차", level=1)
    for sec in toc.get("toc", []):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{sec['order']}. {sec['title']}")
        run.bold = True
        for sub in sec.get("subsections", []):
            sp = doc.add_paragraph(style="List Bullet")
            sp.add_run(f"  {sub}").font.size = Pt(9)

    doc.add_page_break()

    # ── 섹션 초안 ────────────────────────────────────────────────────────────
    doc.add_heading("제안 내용", level=1)
    for sec in toc.get("toc", []):
        title = sec["title"]
        content = sections_draft.get(title, "")
        doc.add_heading(f"{sec['order']}. {title}", level=2)
        if content:
            _md_to_para(doc, content)
        doc.add_page_break()

    # ── 요구사항 대응표 ───────────────────────────────────────────────────────
    rows = matrix.get("matrix", [])
    if rows:
        doc.add_heading("요구사항 대응표", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["요구사항", "대응 섹션", "충족 여부", "비고"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(hdr[i], "1F497D")
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        COLOR_MAP = {"충족": "E6F4EA", "부분충족": "FFF8E1", "미충족": "FCE8E6"}
        for r in rows:
            row = table.add_row().cells
            row[0].text = r.get("requirement", "")
            row[1].text = r.get("section", "")
            row[2].text = r.get("coverage", "")
            row[3].text = r.get("note", "")
            bg = COLOR_MAP.get(r.get("coverage", ""), "FFFFFF")
            for cell in row:
                _set_cell_bg(cell, bg)

    # ── 평가배점표 ────────────────────────────────────────────────────────────
    scores = rfp_result.get("evaluation_scores", [])
    if scores:
        doc.add_page_break()
        doc.add_heading("평가배점표", level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["평가 항목", "배점"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(hdr[i], "1F497D")
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for s in scores:
            row = table.add_row().cells
            row[0].text = s.get("item", "")
            row[1].text = str(s.get("score", ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
